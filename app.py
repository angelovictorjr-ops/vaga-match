"""
MVP — Analisador de Compatibilidade de Currículos com Vagas

Como executar:
1. Crie uma pasta para o projeto.
2. Salve este arquivo como app.py.
3. Salve o conteúdo de requirements.txt no mesmo diretório.
4. Instale as dependências:
       pip install -r requirements.txt
5. Execute:
       streamlit run app.py

Adzuna:
- Para usar a API real, preencha ADZUNA_APP_ID e ADZUNA_APP_KEY abaixo.
- Sem credenciais, o sistema entra automaticamente no MODO DEMONSTRAÇÃO
  e cria 30 vagas fictícias realistas para a região informada.

Observação:
- A pontuação usa TF-IDF + similaridade por cosseno.
- A "análise em tempo real" neste MVP é executada sob demanda quando
  o usuário clica em "Analisar Agora".
"""

import io
import random
import time
from typing import Optional

import pandas as pd
import requests
import streamlit as st
from docx import Document
from PyPDF2 import PdfReader
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


# ============================================================
# CONFIGURAÇÃO
# ============================================================

ADZUNA_APP_ID = ""   # Cole aqui o APP ID do Adzuna
ADZUNA_APP_KEY = ""  # Cole aqui o APP KEY do Adzuna

ADZUNA_COUNTRY = "br"  # Altere para o país desejado quando aplicável
REQUEST_TIMEOUT = 15

st.set_page_config(
    page_title="CV Match AI",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded",
)


# ============================================================
# ESTILO
# ============================================================

st.markdown(
    """
    <style>
    .main {
        background: #f7f9fc;
    }

    .block-container {
        max-width: 1200px;
        padding-top: 2rem;
    }

    .hero {
        padding: 1.4rem 1.6rem;
        border-radius: 18px;
        background: linear-gradient(135deg, #0f172a, #1e3a8a);
        color: white;
        margin-bottom: 1.5rem;
    }

    .hero h1 {
        margin: 0;
        font-size: 2.2rem;
    }

    .hero p {
        margin: .5rem 0 0 0;
        opacity: .9;
    }

    .job-card {
        background: white;
        border: 1px solid #e5e7eb;
        border-radius: 16px;
        padding: 1.2rem;
        margin: .8rem 0;
        box-shadow: 0 4px 16px rgba(15, 23, 42, .06);
    }

    .job-title {
        font-size: 1.15rem;
        font-weight: 700;
        color: #111827;
    }

    .job-company {
        color: #4b5563;
        margin-top: .2rem;
    }

    .job-location {
        color: #6b7280;
        font-size: .9rem;
        margin-top: .35rem;
    }

    .score {
        font-size: 1.8rem;
        font-weight: 800;
        color: #2563eb;
    }

    .badge {
        display: inline-block;
        padding: .25rem .55rem;
        border-radius: 999px;
        font-size: .75rem;
        font-weight: 700;
        background: #dbeafe;
        color: #1d4ed8;
        margin-bottom: .4rem;
    }

    .info-box {
        padding: .9rem 1rem;
        border-radius: 12px;
        background: #eff6ff;
        border: 1px solid #bfdbfe;
        color: #1e3a8a;
    }

    .demo-box {
        padding: .9rem 1rem;
        border-radius: 12px;
        background: #fffbeb;
        border: 1px solid #fde68a;
        color: #92400e;
    }
    </style>
    """,
    unsafe_allow_html=True,
)


# ============================================================
# EXTRAÇÃO DE CURRÍCULO
# ============================================================

@st.cache_data(show_spinner=False)
def extract_resume_text(file_bytes: bytes, file_name: str) -> str:
    """Extrai texto de PDF ou DOCX e guarda o resultado em cache."""
    extension = file_name.lower().rsplit(".", 1)[-1]

    if extension == "pdf":
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []

        for page in reader.pages:
            pages.append(page.extract_text() or "")

        text = "\n".join(pages)

    elif extension == "docx":
        document = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in document.paragraphs]

        # Também tenta extrair texto de tabelas.
        for table in document.tables:
            for row in table.rows:
                paragraphs.append(" | ".join(cell.text for cell in row.cells))

        text = "\n".join(paragraphs)

    else:
        raise ValueError("Formato não suportado. Use PDF ou DOCX.")

    text = " ".join(text.split())

    if not text:
        raise ValueError(
            "Não foi possível extrair texto do currículo. "
            "Verifique se o PDF possui texto selecionável."
        )

    return text


# ============================================================
# ADZUNA
# ============================================================

def adzuna_configured() -> bool:
    return bool(ADZUNA_APP_ID.strip() and ADZUNA_APP_KEY.strip())


def fetch_adzuna_jobs(city: str, state: str, limit: int = 30) -> pd.DataFrame:
    """
    Consulta a API do Adzuna.

    O endpoint utiliza o país configurado em ADZUNA_COUNTRY.
    A localização é enviada como cidade + estado.
    """
    location = f"{city}, {state}".strip(", ")

    url = (
        f"https://api.adzuna.com/v1/api/jobs/"
        f"{ADZUNA_COUNTRY}/search/1"
    )

    params = {
        "app_id": ADZUNA_APP_ID,
        "app_key": ADZUNA_APP_KEY,
        "results_per_page": min(limit, 50),
        "where": location,
        "content-type": "application/json",
    }

    response = requests.get(
        url,
        params=params,
        timeout=REQUEST_TIMEOUT,
    )
    response.raise_for_status()

    data = response.json()
    results = data.get("results", [])

    rows = []

    for job in results:
        rows.append(
            {
                "title": job.get("title", "Vaga sem título"),
                "company": (job.get("company") or {}).get(
                    "display_name", "Empresa não informada"
                ),
                "location": (job.get("location") or {}).get(
                    "display_name", location
                ),
                "description": job.get("description", ""),
                "url": job.get("redirect_url", "#"),
                "source": "Adzuna",
            }
        )

    return pd.DataFrame(rows)


# ============================================================
# MODO DEMONSTRAÇÃO
# ============================================================

def generate_demo_jobs(city: str, state: str, count: int = 30) -> pd.DataFrame:
    """Gera vagas fictícias, porém realistas, para testes do MVP."""

    templates = [
        (
            "Analista de Dados",
            "DataCorp",
            "Python SQL Power BI Excel análise de dados estatística dashboards "
            "ETL banco de dados indicadores KPI",
        ),
        (
            "Desenvolvedor Python",
            "Tech Solutions",
            "Python Django Flask APIs REST PostgreSQL Git Docker testes "
            "desenvolvimento backend orientação a objetos",
        ),
        (
            "Desenvolvedor Java",
            "SoftWay",
            "Java Spring Boot APIs REST SQL PostgreSQL Git Docker microsserviços "
            "desenvolvimento backend",
        ),
        (
            "Desenvolvedor Mobile Flutter",
            "Mobile Labs",
            "Flutter Dart Android iOS Firebase REST API Git UI responsiva "
            "desenvolvimento mobile",
        ),
        (
            "Administrador de Redes",
            "NetWorks",
            "TCP IP VLAN VPN Cisco Linux Windows servidores firewall redes "
            "monitoramento infraestrutura",
        ),
        (
            "Analista de Sistemas",
            "Digital Systems",
            "levantamento de requisitos UML SQL sistemas integração APIs "
            "documentação testes análise de processos",
        ),
        (
            "Engenheiro de Dados Júnior",
            "DataFlow",
            "Python SQL ETL pipelines Airflow PostgreSQL AWS dados cloud "
            "modelagem de dados Spark",
        ),
        (
            "Suporte Técnico",
            "HelpTech",
            "Windows Linux hardware redes atendimento troubleshooting suporte "
            "técnico manutenção informática",
        ),
        (
            "Desenvolvedor Full Stack",
            "WebFactory",
            "Python JavaScript React HTML CSS APIs REST PostgreSQL Git Docker "
            "frontend backend full stack",
        ),
        (
            "Analista de BI",
            "Business Intelligence Pro",
            "Power BI SQL Excel dashboards DAX indicadores relatórios ETL "
            "análise de negócio",
        ),
    ]

    companies = [
        "NovaTech",
        "Moz Digital",
        "Global Systems",
        "Smart Solutions",
        "Prime IT",
        "Inova Labs",
    ]

    random.seed(f"{city}-{state}-cv-match-demo")

    rows = []

    for index in range(count):
        title, default_company, skills = templates[index % len(templates)]

        # Pequena variação para parecer uma lista realista.
        company = (
            default_company
            if index < len(templates)
            else random.choice(companies)
        )

        rows.append(
            {
                "title": title,
                "company": company,
                "location": f"{city}, {state}",
                "description": (
                    f"Estamos buscando {title}. "
                    f"Atuação em projetos de tecnologia. "
                    f"Requisitos e competências: {skills}. "
                    f"Experiência com trabalho em equipe, resolução de problemas, "
                    f"comunicação e aprendizado contínuo."
                ),
                "url": (
                    "https://www.google.com/search?q="
                    + requests.utils.quote(
                        f"{title} {city} {state} vaga"
                    )
                ),
                "source": "Demonstração",
            }
        )

    return pd.DataFrame(rows)


def get_jobs(city: str, state: str) -> tuple[pd.DataFrame, bool, Optional[str]]:
    """
    Retorna vagas, indicador de modo demonstração e mensagem de erro da API.
    """
    if not adzuna_configured():
        return generate_demo_jobs(city, state, 30), True, None

    try:
        jobs = fetch_adzuna_jobs(city, state, 30)

        if jobs.empty:
            return (
                generate_demo_jobs(city, state, 30),
                True,
                "A API não retornou vagas para essa região.",
            )

        return jobs, False, None

    except requests.RequestException as exc:
        return (
            generate_demo_jobs(city, state, 30),
            True,
            f"Falha na API do Adzuna: {exc}",
        )

    except Exception as exc:
        return (
            generate_demo_jobs(city, state, 30),
            True,
            f"Erro ao consultar o Adzuna: {exc}",
        )


# ============================================================
# MATCHING
# ============================================================

def calculate_match_scores(
    resume_text: str,
    jobs: pd.DataFrame,
) -> pd.DataFrame:
    """
    Calcula TF-IDF entre currículo e descrição de cada vaga
    e transforma a similaridade em uma escala de 0 a 100.
    """
    if jobs.empty:
        return jobs.copy()

    job_texts = (
        jobs["title"].fillna("")
        + " "
        + jobs["description"].fillna("")
    ).tolist()

    documents = [resume_text] + job_texts

    vectorizer = TfidfVectorizer(
        lowercase=True,
        strip_accents="unicode",
        stop_words=None,
        ngram_range=(1, 2),
        max_features=12000,
    )

    matrix = vectorizer.fit_transform(documents)

    similarities = cosine_similarity(
        matrix[0:1],
        matrix[1:],
    )[0]

    result = jobs.copy()
    result["score"] = (similarities * 100).clip(0, 100).round(1)
    result = result.sort_values(
        by="score",
        ascending=False,
    ).reset_index(drop=True)

    return result


# ============================================================
# UI AUXILIAR
# ============================================================

def score_color(score: float) -> str:
    if score >= 75:
        return "green"
    if score >= 50:
        return "orange"
    return "red"


def render_job_card(job: pd.Series, rank: int) -> None:
    score = float(job["score"])
    source = job.get("source", "Vaga")

    st.markdown('<div class="job-card">', unsafe_allow_html=True)

    left, right = st.columns([4, 1])

    with left:
        st.markdown(
            f'<div class="badge">#{rank} • {source}</div>',
            unsafe_allow_html=True,
        )
        st.markdown(
            f'<div class="job-title">{job["title"]}</div>',
            unsafe_allow_html=True,
        )
        st.markdown(
            f'<div class="job-company">{job["company"]}</div>',
            unsafe_allow_html=True,
        )
        st.markdown(
            f'<div class="job-location">📍 {job["location"]}</div>',
            unsafe_allow_html=True,
        )

    with right:
        st.markdown(
            f'<div class="score">{score:.1f}%</div>',
            unsafe_allow_html=True,
        )

    st.progress(
        int(round(score)),
        text=f"Compatibilidade: {score:.1f}%",
    )

    button_col, details_col = st.columns([1, 3])

    with button_col:
        url = str(job.get("url", "#"))

        if url and url != "#":
            st.link_button(
                "Candidatar-se",
                url,
                use_container_width=True,
            )
        else:
            st.button(
                "Link indisponível",
                disabled=True,
                use_container_width=True,
            )

    with details_col:
        with st.expander("Ver descrição da vaga"):
            st.write(job["description"])

    st.markdown("</div>", unsafe_allow_html=True)


# ============================================================
# APLICAÇÃO
# ============================================================

def main() -> None:
    st.markdown(
        """
        <div class="hero">
            <h1>🎯 CV Match AI</h1>
            <p>
                Compare seu currículo com vagas da sua região usando
                TF-IDF + similaridade por cosseno.
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # -----------------------------
    # SIDEBAR
    # -----------------------------
    with st.sidebar:
        st.header("📍 Região")

        city = st.text_input(
            "Cidade",
            placeholder="Ex.: Maputo",
        )

        state = st.text_input(
            "Estado / Província",
            placeholder="Ex.: Maputo",
        )

        st.divider()

        st.header("📄 Currículo")

        uploaded_file = st.file_uploader(
            "Envie seu currículo",
            type=["pdf", "docx"],
            help="Formatos aceitos: PDF e DOCX.",
        )

        st.divider()

        if adzuna_configured():
            st.success("Adzuna configurado")
        else:
            st.warning("Modo Demonstração ativo")

        st.caption(
            "Sem APP_ID e APP_KEY, o sistema usa 30 vagas fictícias "
            "para permitir testar o algoritmo."
        )

    # -----------------------------
    # INFORMAÇÕES INICIAIS
    # -----------------------------
    if not city or not state:
        st.info(
            "Preencha Cidade e Estado/Província na barra lateral "
            "para começar."
        )

    if not uploaded_file:
        st.info(
            "Envie um currículo em PDF ou DOCX para habilitar a análise."
        )

    if city and state and uploaded_file:
        st.success(
            f"Currículo carregado: **{uploaded_file.name}** — "
            f"região selecionada: **{city}, {state}**"
        )

        analyze = st.button(
            "🚀 Analisar Agora",
            type="primary",
            use_container_width=True,
        )

        if analyze:
            try:
                with st.spinner(
                    "Extraindo currículo, buscando vagas e calculando "
                    "compatibilidade..."
                ):
                    # Simulação de processamento para UX do MVP.
                    time.sleep(1.5)

                    file_bytes = uploaded_file.getvalue()

                    resume_text = extract_resume_text(
                        file_bytes,
                        uploaded_file.name,
                    )

                    time.sleep(1.0)

                    jobs, demo_mode, api_error = get_jobs(
                        city,
                        state,
                    )

                    time.sleep(1.0)

                    results = calculate_match_scores(
                        resume_text,
                        jobs,
                    )

                    time.sleep(1.0)

                    st.session_state["results"] = results
                    st.session_state["resume_text"] = resume_text
                    st.session_state["demo_mode"] = demo_mode
                    st.session_state["api_error"] = api_error
                    st.session_state["region"] = f"{city}, {state}"

            except Exception as exc:
                st.error(f"Não foi possível analisar o currículo: {exc}")

    # -----------------------------
    # RESULTADOS
    # -----------------------------
    results = st.session_state.get("results")

    if results is not None:
        st.divider()

        st.subheader("📊 Resultado da análise")

        demo_mode = st.session_state.get("demo_mode", True)
        api_error = st.session_state.get("api_error")

        if demo_mode:
            st.markdown(
                """
                <div class="demo-box">
                    <strong>Modo Demonstração:</strong>
                    as vagas exibidas são fictícias e foram geradas para
                    testar o algoritmo de compatibilidade.
                </div>
                """,
                unsafe_allow_html=True,
            )

            if api_error:
                st.warning(api_error)
        else:
            st.markdown(
                """
                <div class="info-box">
                    <strong>Fonte:</strong>
                    vagas obtidas através da API do Adzuna.
                </div>
                """,
                unsafe_allow_html=True,
            )

        st.write("")

        avg_score = float(results["score"].mean()) if not results.empty else 0
        best_score = float(results["score"].max()) if not results.empty else 0

        metric1, metric2, metric3 = st.columns(3)

        with metric1:
            st.metric(
                "Vagas analisadas",
                len(results),
            )

        with metric2:
            st.metric(
                "Melhor compatibilidade",
                f"{best_score:.1f}%",
            )

        with metric3:
            st.metric(
                "Média de compatibilidade",
                f"{avg_score:.1f}%",
            )

        st.write("")
        st.subheader("🏆 Vagas mais compatíveis")

        for index, (_, job) in enumerate(results.iterrows(), start=1):
            render_job_card(job, index)


if __name__ == "__main__":
    main()
